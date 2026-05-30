from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import os

class ReporteFinal:
    def generar_reporte(self, simulador, analisis_sensibilidad=None):
        doc = Document()
        doc.add_heading('Análisis Epidemiológico y Económico de Trabajo Sexual por Necesidad', 0)
        doc.add_paragraph(f'Fecha: {__import__("datetime").datetime.now().strftime("%Y-%m-%d")}')

        # Resumen ejecutivo
        doc.add_heading('Resumen Ejecutivo', level=1)
        doc.add_paragraph('Modelo agent-based que diferencia prostitución callejera vs VIP, incorpora violencia, costos reales y feedback económico.')

        # Resultados
        doc.add_heading('Resultados Principales', level=1)
        doc.add_paragraph(f"Mujeres simuladas: {len(simulador.mujeres)}")
        doc.add_paragraph(f"Prevalencia VIH final: {simulador.prevalencia['VIH'][-1]*100:.2f}%")

        # Gráficos
        doc.add_heading('Gráficos', level=1)
        # (guardar gráficos temporalmente)
        plt.figure(figsize=(10,6))
        for pat in simulador.prevalencia:
            plt.plot(simulador.prevalencia[pat], label=pat)
        plt.legend()
        plt.savefig("outputs/temp_prevalencia.png")
        doc.add_picture("outputs/temp_prevalencia.png", width=Inches(6))
        
        doc.save("outputs/REPORTE_FINAL_COMPLETO.docx")
        print("✅ Reporte Word generado: outputs/REPORTE_FINAL_COMPLETO.docx")